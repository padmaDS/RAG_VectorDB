
## Extracting all information from an image and saving it in a docx file


import base64
import requests
from openai import OpenAI
import os
from dotenv import load_dotenv
import json
from docx import Document

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI()

# Function to encode the image
def encode_image(image_path):
  with open(image_path, "rb") as image_file:
    return base64.b64encode(image_file.read()).decode('utf-8')

# Path to your image
image_path = "images\Payroll_Policy_page-0007.jpg"

# Getting the base64 string
base64_image = encode_image(image_path)

headers = {
  "Content-Type": "application/json",
  "Authorization": f"Bearer {api_key}"
}

payload = {
  "model": "gpt-4o",
  "messages": [
    {
      "role": "user",
      "content": [
        {
          "type": "text",
          "text": "extract all the text from this image and store in a variable?"
        },
        {
          "type": "image_url",
          "image_url": {
            "url": f"data:image/jpeg;base64,{base64_image}"
          }
        }
      ]
    }
  ],
  "max_tokens": 1000
}

response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)

# print(response.json())

response_data = json.loads(response.content)

# Extract relevant details
message_content = response_data['choices'][0]['message']['content']

# Print relevant details
print(message_content)

## Storing in a docx file

doc = Document()
doc.add_paragraph(message_content)
doc.save('Payroll_Policy.docx')

### Open existing DOCX file
existing_docx_path = "Payroll_Policy.docx"
doc = Document(existing_docx_path)

# Add extracted text to the existing document
doc.add_paragraph(message_content)

# Save the modified document
doc.save(existing_docx_path)